#!/usr/bin/env python3
"""Build Phase 0 data files from the Appendix 2 DOCX table."""

from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
DOCX_PATH = ROOT / "appendix 2 附錄_明洪武時期蕃國朝貢一覽表(6.1.2024).docx"
OUT_DIR = ROOT / "tribute-map" / "js"

CHINESE_NUMBERS = {
    "零": 0,
    "〇": 0,
    "一": 1,
    "二": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
}

COUNTRIES = {
    "goryeo": {
        "id": "goryeo",
        "name": "高丽",
        "nameEn": "Goryeo",
        "lat": 37.97,
        "lng": 126.56,
        "region": "korea",
        "routeType": "land_north",
        "note": "洪武二十五年前主要以高丽名义出现。",
    },
    "joseon": {
        "id": "joseon",
        "name": "朝鲜",
        "nameEn": "Joseon",
        "lat": 37.57,
        "lng": 126.98,
        "region": "korea",
        "routeType": "land_north",
        "note": "洪武二十六年起正式使用朝鲜国号；高丽（朝鲜）记录按朝鲜处理。",
    },
    "annam": {
        "id": "annam",
        "name": "安南",
        "nameEn": "Annam",
        "lat": 21.03,
        "lng": 105.85,
        "region": "southeast_asia",
        "routeType": "sea_south",
        "note": "以升龙/东都，即今河内一带定位。",
    },
    "champa": {
        "id": "champa",
        "name": "占城",
        "nameEn": "Champa",
        "lat": 15.35,
        "lng": 108.8,
        "region": "southeast_asia",
        "routeType": "sea_south",
        "note": "以占城核心区近今越南中南部定位。",
    },
    "siam": {
        "id": "siam",
        "name": "暹罗",
        "nameEn": "Siam",
        "lat": 14.35,
        "lng": 100.52,
        "region": "southeast_asia",
        "routeType": "sea_south",
        "note": "以阿瑜陀耶，即今大城一带定位。",
    },
    "cambodia": {
        "id": "cambodia",
        "name": "真腊",
        "nameEn": "Chenla / Cambodia",
        "lat": 13.36,
        "lng": 103.86,
        "region": "southeast_asia",
        "routeType": "sea_south",
        "note": "以吴哥，即今暹粒一带定位。",
    },
    "java": {
        "id": "java",
        "name": "爪哇",
        "nameEn": "Java",
        "lat": -7.28,
        "lng": 112.75,
        "region": "island_southeast_asia",
        "routeType": "sea_south",
        "note": "以满者伯夷核心区近今泗水一带定位。",
    },
    "java_east": {
        "id": "java_east",
        "name": "爪哇东番",
        "nameEn": "Eastern Java polity",
        "lat": -7.28,
        "lng": 112.95,
        "region": "island_southeast_asia",
        "routeType": "sea_south",
        "note": "洪武十二年一次记录；与爪哇共用核心区域，略作视觉偏移。",
    },
    "java_west": {
        "id": "java_west",
        "name": "爪哇西番",
        "nameEn": "Western Java polity",
        "lat": -7.28,
        "lng": 112.55,
        "region": "island_southeast_asia",
        "routeType": "sea_south",
        "note": "洪武十二年一次记录；与爪哇共用核心区域，略作视觉偏移。",
    },
    "srivijaya": {
        "id": "srivijaya",
        "name": "三佛齐",
        "nameEn": "Srivijaya",
        "lat": -2.99,
        "lng": 104.75,
        "region": "island_southeast_asia",
        "routeType": "west_sea",
        "note": "以旧港/巨港一带定位。",
    },
    "japan": {
        "id": "japan",
        "name": "日本",
        "nameEn": "Japan",
        "lat": 33.0,
        "lng": 130.5,
        "region": "east_asia",
        "routeType": "east_sea",
        "note": "以洪武时期主要对华港博多一带定位。",
    },
    "ryukyu_chuzan": {
        "id": "ryukyu_chuzan",
        "name": "琉球中山国",
        "nameEn": "Ryukyu Chuzan",
        "lat": 26.33,
        "lng": 127.77,
        "region": "east_asia",
        "routeType": "east_sea",
        "note": "泛称“琉球”的洪武十八年记录归入中山国。",
    },
    "ryukyu_sannan": {
        "id": "ryukyu_sannan",
        "name": "琉球山南国",
        "nameEn": "Ryukyu Sannan",
        "lat": 26.15,
        "lng": 127.7,
        "region": "east_asia",
        "routeType": "east_sea",
        "note": "琉球三国之一，独立显示。",
    },
    "ryukyu_hokuzan": {
        "id": "ryukyu_hokuzan",
        "name": "琉球山北国",
        "nameEn": "Ryukyu Hokuzan",
        "lat": 26.7,
        "lng": 127.9,
        "region": "east_asia",
        "routeType": "east_sea",
        "note": "琉球三国之一，独立显示。",
    },
    "brunei": {
        "id": "brunei",
        "name": "浡泥",
        "nameEn": "Brunei",
        "lat": 4.94,
        "lng": 114.95,
        "region": "island_southeast_asia",
        "routeType": "sea_south",
        "note": "以今文莱一带定位。",
    },
    "xiyang": {
        "id": "xiyang",
        "name": "西洋",
        "nameEn": "Xiyang",
        "lat": 11.25,
        "lng": 75.78,
        "region": "indian_ocean",
        "routeType": "west_sea",
        "note": "古地名对应仍需学术确认，暂按南印度西岸方向定位。",
    },
    "soli": {
        "id": "soli",
        "name": "琐里",
        "nameEn": "Soli",
        "lat": 10.8,
        "lng": 79.0,
        "region": "indian_ocean",
        "routeType": "west_sea",
        "note": "古地名对应仍需学术确认，暂按南印度方向定位。",
    },
    "fulin": {
        "id": "fulin",
        "name": "拂菻国",
        "nameEn": "Fulin",
        "lat": 41.01,
        "lng": 75.2,
        "actualLat": 41.01,
        "actualLng": 28.98,
        "region": "far_west",
        "routeType": "west_sea",
        "positionStatus": "edge_marker",
        "note": "若按拜占庭方向理解，实际位置远超主图范围；主图使用西侧边界方向标记。",
    },
    "lambri": {
        "id": "lambri",
        "name": "览邦",
        "nameEn": "Lambri",
        "lat": 5.55,
        "lng": 95.32,
        "region": "indian_ocean",
        "routeType": "west_sea",
        "note": "以苏门答腊北端兰无里一带定位。",
    },
    "timor": {
        "id": "timor",
        "name": "淡巴",
        "nameEn": "Tamba",
        "lat": 2.0,
        "lng": 103.0,
        "region": "southeast_asia",
        "routeType": "west_sea",
        "note": "古地名对应仍需学术确认。",
    },
    "pahang": {
        "id": "pahang",
        "name": "湓亨",
        "nameEn": "Pahang",
        "lat": 3.8,
        "lng": 103.4,
        "region": "southeast_asia",
        "routeType": "west_sea",
        "note": "以今马来西亚彭亨州一带定位。",
    },
    "baihua": {
        "id": "baihua",
        "name": "百花",
        "nameEn": "Baihua",
        "lat": -7.8,
        "lng": 110.36,
        "region": "island_southeast_asia",
        "routeType": "sea_south",
        "note": "古地名对应仍需学术确认，暂按爪哇中部方向定位。",
    },
    "jawa": {
        "id": "jawa",
        "name": "阇婆",
        "nameEn": "Jawa",
        "lat": -7.6,
        "lng": 110.4,
        "region": "island_southeast_asia",
        "routeType": "sea_south",
        "note": "古地名对应仍需学术确认，暂按爪哇方向定位。",
    },
    "samudra": {
        "id": "samudra",
        "name": "须文那达",
        "nameEn": "Samudra",
        "lat": 5.18,
        "lng": 97.14,
        "region": "indian_ocean",
        "routeType": "west_sea",
        "note": "以苏门答腊/巴赛方向定位。",
    },
}

COUNTRY_NAME_TO_ID = {
    "高丽": "goryeo",
    "朝鲜": "joseon",
    "高丽（朝鲜）": "joseon",
    "安南": "annam",
    "占城": "champa",
    "暹罗": "siam",
    "真腊": "cambodia",
    "爪哇": "java",
    "爪哇东番": "java_east",
    "爪哇西番": "java_west",
    "三佛齐": "srivijaya",
    "日本": "japan",
    "琉球": "ryukyu_chuzan",
    "琉球中山国": "ryukyu_chuzan",
    "琉球山南国": "ryukyu_sannan",
    "琉球山北国": "ryukyu_hokuzan",
    "浡泥": "brunei",
    "西洋": "xiyang",
    "琐里": "soli",
    "拂菻国": "fulin",
    "览邦": "lambri",
    "淡巴": "timor",
    "湓亨": "pahang",
    "百花": "baihua",
    "阇婆": "jawa",
    "须文那达": "samudra",
}

CATEGORY_KEYWORDS = {
    "animal": [
        "虎",
        "象",
        "马",
        "鹿",
        "熊",
        "猴",
        "孔雀",
        "鹦鹉",
        "火鸡",
        "玳瑁",
        "龟",
        "犀",
        "番奴",
        "侍童",
        "阉",
    ],
    "spice": [
        "胡椒",
        "苏木",
        "丁香",
        "香",
        "檀",
        "降",
        "木香",
        "没药",
        "龙脑",
        "米脑",
        "乳香",
        "豆蔻",
        "蔷薇",
        "槟榔",
        "波罗蜜",
    ],
    "textile": [
        "布",
        "帛",
        "罗",
        "绮",
        "纱",
        "绢",
        "绵",
        "被",
        "衣",
        "苾布",
        "疋",
        "匹",
    ],
    "vessel": [
        "盏",
        "盘",
        "樽",
        "器",
        "壶",
        "罐",
        "钟",
        "盆",
        "酒",
        "鞍",
        "刀",
        "甲",
        "印",
    ],
}

ROUTE_TEMPLATES = {
    "sea_south": {
        "type": "sea_south",
        "name": "东南亚海路",
        "waypoints": [
            {"lat": 8.0, "lng": 109.0},
            {"lat": 18.5, "lng": 112.0},
            {"lat": 24.8, "lng": 118.6},
            {"lat": 30.0, "lng": 121.2},
        ],
    },
    "east_sea": {
        "type": "east_sea",
        "name": "东海路线",
        "waypoints": [
            {"lat": 27.0, "lng": 124.2},
            {"lat": 29.5, "lng": 122.0},
            {"lat": 31.0, "lng": 120.4},
        ],
    },
    "land_north": {
        "type": "land_north",
        "name": "陆路/混合路线",
        "waypoints": [
            {"lat": 40.0, "lng": 124.4},
            {"lat": 38.0, "lng": 121.5},
            {"lat": 35.5, "lng": 119.0},
        ],
    },
    "west_sea": {
        "type": "west_sea",
        "name": "西域/印度洋路线",
        "waypoints": [
            {"lat": 5.0, "lng": 95.0},
            {"lat": 9.0, "lng": 104.0},
            {"lat": 18.5, "lng": 112.0},
            {"lat": 24.8, "lng": 118.6},
            {"lat": 30.0, "lng": 121.2},
        ],
    },
}


def chinese_to_int(text: str) -> int:
    text = text.strip()
    if text in ("正", "元"):
        return 1
    if text in CHINESE_NUMBERS:
        return CHINESE_NUMBERS[text]
    if text == "十":
        return 10
    if "十" in text:
        left, right = text.split("十", 1)
        tens = CHINESE_NUMBERS.get(left, 1) if left else 1
        ones = CHINESE_NUMBERS.get(right, 0) if right else 0
        return tens * 10 + ones
    raise ValueError(f"Cannot parse Chinese number: {text}")


def parse_time(value: str) -> dict:
    match = re.search(r"洪武(.+?)年(闰)?(.+?)月", value)
    if not match:
        raise ValueError(f"Cannot parse time field: {value}")
    year_text, leap_text, month_text = match.groups()
    year = chinese_to_int(year_text)
    month = chinese_to_int(month_text)
    is_leap_month = bool(leap_text)
    return {
        "year": year,
        "month": month,
        "isLeapMonth": is_leap_month,
        "monthSort": month + (0.5 if is_leap_month else 0),
        "gregorianYear": year + 1367,
    }


def clean_cell(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\n", " / ")).strip()


def normalize_item(value: str) -> str:
    value = re.sub(r"\([^)]*\)|（[^）]*）", "", value)
    value = re.sub(r"[0-9一二三四五六七八九十百千万余兩两斤匹疋只足枝副兠颗个人口艘锭]+", "", value)
    value = value.strip(" ：:。；;、，, /")
    return value


def split_tribute_items(value: str) -> list[str]:
    if not value:
        return []
    parts = re.split(r"[、，,；;。/]|并|及|等物|之属|凡|又贡|合共计", value)
    items = []
    generic_terms = {
        "方物",
        "本国地图",
        "贡中宫东宫",
        "使臣亦自有献",
        "与上同一项",
        "与上项",
        "与上项合共计",
        "与上同项",
    }
    for part in parts:
        item = normalize_item(part)
        if item and item not in generic_terms:
            items.append(item)
    return items


def categorize_item(item: str) -> str:
    for category, keywords in CATEGORY_KEYWORDS.items():
        if any(keyword in item for keyword in keywords):
            return category
    return "misc"


def make_record_id(country_id: str, time_info: dict, ordinal: int) -> str:
    leap = "r" if time_info["isLeapMonth"] else ""
    return f"{country_id}_{time_info['year']:02d}_{leap}{time_info['month']:02d}_{ordinal:03d}"


def read_records() -> list[dict]:
    document = Document(DOCX_PATH)
    if len(document.tables) != 1:
        raise ValueError(f"Expected exactly one table, found {len(document.tables)}")

    rows = document.tables[0].rows
    records = []
    id_counts: Counter[tuple[str, int, int, bool]] = Counter()
    for row_index, row in enumerate(rows[1:], start=2):
        cells = [clean_cell(cell.text) for cell in row.cells]
        time_raw, country_raw, king, envoy, purpose, tribute_raw, gifts_raw, source = cells[1:9]
        if not time_raw or not country_raw:
            raise ValueError(f"Missing required time/country at table row {row_index}")

        country_id = COUNTRY_NAME_TO_ID.get(country_raw)
        if not country_id:
            raise ValueError(f"Unknown country name at table row {row_index}: {country_raw}")

        time_info = parse_time(time_raw)
        id_key = (country_id, time_info["year"], time_info["month"], time_info["isLeapMonth"])
        id_counts[id_key] += 1

        tribute_items = split_tribute_items(tribute_raw)
        tribute_categories = sorted({categorize_item(item) for item in tribute_items}) or ["misc"]
        country_name = COUNTRIES[country_id]["name"]

        record = {
            "id": make_record_id(country_id, time_info, id_counts[id_key]),
            "sourceRow": row_index,
            "countryId": country_id,
            "countryName": country_name,
            "countryRaw": country_raw,
            "timeRaw": time_raw,
            **time_info,
            "king": king,
            "envoy": envoy,
            "purpose": purpose,
            "tributeRaw": tribute_raw,
            "tributeItems": tribute_items,
            "tributeCategories": tribute_categories,
            "giftsRaw": gifts_raw,
            "source": source,
        }
        records.append(record)
    return records


def enrich_countries(records: list[dict]) -> dict:
    enriched = json.loads(json.dumps(COUNTRIES, ensure_ascii=False))
    counts = Counter(record["countryId"] for record in records)
    first_year = {}
    last_year = {}
    for record in records:
        cid = record["countryId"]
        first_year[cid] = min(first_year.get(cid, record["year"]), record["year"])
        last_year[cid] = max(last_year.get(cid, record["year"]), record["year"])

    for cid, country in enriched.items():
        country["totalRecords"] = counts.get(cid, 0)
        country["firstYear"] = first_year.get(cid)
        country["lastYear"] = last_year.get(cid)
    return enriched


def write_js_file(path: Path, constants: dict, body: str = "") -> None:
    chunks = [
        "/* Auto-generated by tribute-map/scripts/build_phase0_data.py. Do not edit by hand. */",
        "(function (global) {",
        '  "use strict";',
    ]
    for name, value in constants.items():
        rendered = json.dumps(value, ensure_ascii=False, indent=2)
        chunks.append(f"  const {name} = {rendered};")
        chunks.append(f"  global.{name} = {name};")
    if body:
        chunks.append(body)
    export_names = ", ".join(constants.keys())
    chunks.append("  if (typeof module !== \"undefined\" && module.exports) {")
    chunks.append(f"    module.exports = {{ {export_names} }};")
    chunks.append("  }")
    chunks.append("})(typeof window !== \"undefined\" ? window : globalThis);")
    path.write_text("\n".join(chunks) + "\n", encoding="utf-8")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    records = read_records()
    countries = enrich_countries(records)
    category_counts = Counter(category for record in records for category in record["tributeCategories"])

    data_helpers = """
  global.getRecordsByYear = function getRecordsByYear(year) {
    return TRIBUTE_DATA.filter((record) => record.year === Number(year));
  };
  global.getActiveCountries = function getActiveCountries(year) {
    return new Set(global.getRecordsByYear(year).map((record) => record.countryId));
  };
  global.getRecordsByCountry = function getRecordsByCountry(countryId) {
    return TRIBUTE_DATA.filter((record) => record.countryId === countryId)
      .sort((a, b) => a.year - b.year || a.monthSort - b.monthSort || a.sourceRow - b.sourceRow);
  };
"""

    write_js_file(
        OUT_DIR / "data.js",
        {
            "TRIBUTE_DATA": records,
            "COUNTRIES": countries,
            "COUNTRY_NAME_TO_ID": COUNTRY_NAME_TO_ID,
            "TRIBUTE_META": {
                "recordCount": len(records),
                "sourceDocx": DOCX_PATH.name,
                "sourceTableRows": len(records) + 1,
                "categoryCounts": dict(category_counts),
                "generatedBy": "tribute-map/scripts/build_phase0_data.py",
            },
        },
        data_helpers,
    )

    write_js_file(
        OUT_DIR / "config.js",
        {
            "NANJING": {"name": "应天府", "lat": 32.06, "lng": 118.78},
            "MAP_CONFIG": {
                "bounds": {"latMin": -12, "latMax": 52, "lngMin": 73, "lngMax": 152},
                "projection": "equirectangular",
                "minSupportedWidth": 1024,
            },
            "ROUTE_TEMPLATES": ROUTE_TEMPLATES,
            "TRIBUTE_COLORS": {
                "paper": "#F5F0E8",
                "ink": "#1A1008",
                "vermillion": "#C0392B",
                "goldBrown": "#8B6914",
                "land": "#E8E0D0",
                "coast": "#8B8070",
                "muted": "#B0A898",
            },
            "TRIBUTE_CATEGORIES": {
                "animal": {"label": "动物", "icon": "elephant"},
                "spice": {"label": "香料", "icon": "incense"},
                "textile": {"label": "织物", "icon": "scroll"},
                "vessel": {"label": "器物", "icon": "vase"},
                "misc": {"label": "方物", "icon": "box"},
            },
        },
    )

    print(f"Generated {len(records)} records")
    print(f"Generated {len(countries)} countries")
    print(f"Wrote {OUT_DIR / 'data.js'}")
    print(f"Wrote {OUT_DIR / 'config.js'}")


if __name__ == "__main__":
    main()
